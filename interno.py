from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime, time, date

#Adaptação para uso com pyinstaller
import sys
import os

class InternoGenerator:
    def __init__(self, data_handler, numero_inicial, dia_de_envio, mes_pagamento, responsavel_assinatura):
        self.data_handler = data_handler
        self.numero_inicial = numero_inicial
        self.numero_atual = numero_inicial
        self.dia_de_envio = dia_de_envio
        self.mes_pagamento = mes_pagamento
        self.responsavel_assinatura = responsavel_assinatura
        self.ano_atual = datetime.now().year
        self.dict_mes = {
            1: 'janeiro',
            2: 'fevereiro',
            3: 'março',
            4: 'abril',
            5: 'maio',
            6: 'junho',
            7: 'julho',
            8: 'agosto',
            9: 'setembro',
            10: 'outubro',
            11: 'novembro',
            12: 'dezembro'
        }

    def criar_documento_interno(self):
        #self.documento = Document('modelo.docx')
        #cria o modelo para funcionar com o pyinstaller 
        if getattr(sys, 'frozen', False):
            # Caminho do arquivo dentro do pacote PyInstaller
            bundle_dir = sys._MEIPASS
        else:
            # Caminho do arquivo em desenvolvimento
            bundle_dir = os.path.abspath(os.path.dirname(__file__))

        modelo_path = os.path.join(bundle_dir, 'modelo.docx')
        self.documento = Document(modelo_path)

    def criar_numero_interno(self):
        numero_interno = self.documento.add_paragraph()
        numero_interno.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        numero_interno_run = numero_interno.add_run(f'Interno nº {self.numero_atual}/{self.ano_atual}')
        numero_interno_run.bold = True
        numero_interno_run.font.size = Pt(12)
        espaco = self.documento.add_paragraph()


    def criar_data_interno(self):
        data_interno = self.documento.add_paragraph()
        data_interno.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        mes = self.dict_mes[(self.mes_pagamento % 12) + 1]
        data_interno_run = data_interno.add_run(f'São José do Rio Preto, {self.dia_de_envio} de {mes} de {date.today().year}.')
        data_interno_run.font.size = Pt(12)
        espaco = self.documento.add_paragraph()
        espaco = self.documento.add_paragraph()


    def criar_destinatario_interno(self):
        destinatario_interno = self.documento.add_paragraph()
        destinatario_interno.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        destinatario_interno_run = destinatario_interno.add_run("""À Secretaria Municipal de Administração
A/C Coordenadoria de Pagamento 
Bancada da Educação""")
        destinatario_interno_run.bold = True
        destinatario_interno_run.font.size = Pt(12)
        espaco = self.documento.add_paragraph()
        espaco = self.documento.add_paragraph()

    def criar_assunto_interno(self):
        assunto_interno = self.documento.add_paragraph()
        assunto_interno_run = assunto_interno.add_run("Assunto: Alteração pontual de horário.")
        assunto_interno_run.bold = True
        assunto_interno_run.font.size = Pt(12)
        espaco = self.documento.add_paragraph()
        espaco = self.documento.add_paragraph()


    def criar_corpo_interno(self):
        corpo_interno = self.documento.add_paragraph()
        corpo_interno.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        corpo_interno_run = corpo_interno.add_run("""Encaminham-se as alterações pontuais de horários do servidor abaixo descrito:

Horário Homologado:""")
        corpo_interno_run.font.size = Pt(12)

    def criar_tabela_homologado(self, homologado_filtrado):
        table_homologado = self.documento.add_table(rows=homologado_filtrado.shape[0] + 1, cols=homologado_filtrado.shape[1])
        table_homologado.style = 'Table Grid'  # Add borders to the table

        # Set alignment for each cell in the table
        for row in table_homologado.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Set the header row cells' text and formatting
        header_row = table_homologado.rows[0]
        for i, column_name in enumerate(homologado_filtrado.columns):
            header_cell = header_row.cells[i]
            header_cell.text = column_name
            header_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_cell.bold = True

        # Set the remaining cells' text and formatting
        for row in range(homologado_filtrado.shape[0]):
            for col in range(homologado_filtrado.shape[1]):
                cell = table_homologado.cell(row + 1, col)
                cell.text = str(homologado_filtrado.values[row, col])
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        espaco = self.documento.add_paragraph()


    def criar_horario_alterado_interno(self):
        horario_alterado_interno = self.documento.add_paragraph()
        horario_alterado_interno.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        horario_alterado_interno_run = horario_alterado_interno.add_run("Horário Alterado para:")
        horario_alterado_interno_run.font.size = Pt(12)

    def criar_tabela_alteracoes(self, alteracoes_filtradas):
        table_alteracoes = self.documento.add_table(rows=alteracoes_filtradas.shape[0] + 1, cols=alteracoes_filtradas.shape[1])
        table_alteracoes.style = 'Table Grid'  # Add borders to the table

        # Set alignment for each cell in the table
        for row in table_alteracoes.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Set the header row cells' text and formatting
        header_row = table_alteracoes.rows[0]
        for i, column_name in enumerate(alteracoes_filtradas.columns):
            header_cell = header_row.cells[i]
            header_cell.text = column_name
            header_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_cell.bold = True

        # Set the remaining cells' text and formatting
        for row in range(alteracoes_filtradas.shape[0]):
            for col in range(alteracoes_filtradas.shape[1]):
                cell = table_alteracoes.cell(row + 1, col)
                value = alteracoes_filtradas.iloc[row, col]

                # Format dates to display only day and month
                if isinstance(value, datetime):
                    value = value.strftime('%d/%m')

                # Format times to display as hh:mm
                if isinstance(value, time):
                    value = value.strftime('%H:%M')

                cell.text = str(value)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
        espaco = self.documento.add_paragraph()
        espaco = self.documento.add_paragraph()
        espaco = self.documento.add_paragraph()
        espaco = self.documento.add_paragraph()

   
    def obter_nome_responsavel(self):
        if self.responsavel_assinatura == "Silvia Elaine":
            return "_____________________________\nSilvia Elaine da Silva Ganzela\nDiretora de Escola"
        elif self.responsavel_assinatura == "Anelisa Lopes":
            return "_____________________________\nAnelisa Luciene Lopes Regovich\nCoordenadora Pedagógica"
        else:
            return "_____________________________\nResponsável"

    def criar_responsavel_interno(self):
        responsavel_interno = self.documento.add_paragraph()
        responsavel_interno.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        responsavel_interno_run = responsavel_interno.add_run(self.obter_nome_responsavel())
        responsavel_interno_run.bold = True
        responsavel_interno_run.font.size = Pt(12)
        
          

    def gerar_interno(self, matricula):
        homologado_filtrado = self.data_handler.homologado[self.data_handler.homologado['Mat'] == matricula]
        matricula = homologado_filtrado['Mat'].iloc[0]
        alteracoes_filtradas = self.data_handler.get_alteracoes_filtradas(matricula, self.mes_pagamento)
        self.criar_documento_interno()
        self.criar_numero_interno()
        self.criar_data_interno()
        self.criar_destinatario_interno()
        self.criar_assunto_interno()
        self.criar_corpo_interno()
        self.criar_tabela_homologado(homologado_filtrado)
        self.criar_horario_alterado_interno()
        self.criar_tabela_alteracoes(alteracoes_filtradas)
        self.criar_responsavel_interno()
        self.numero_atual += 1 
        nome_interno = f'Interno nº{self.numero_atual}-{self.ano_atual} - Alteração de horário de {homologado_filtrado.iloc[0]["Nome"]}-{self.dict_mes[self.mes_pagamento]}{self.ano_atual}.docx'
        self.documento.save(f'interno/{nome_interno}')


class CargaGenerator:
    def __init__(self, data_handler,mes_pagamento):
        self.data_handler = data_handler
        self.mes_pagamento = mes_pagamento
        self.ano_atual = datetime.now().year
        self.dict_mes = {
            1: 'janeiro',
            2: 'fevereiro',
            3: 'março',
            4: 'abril',
            5: 'maio',
            6: 'junho',
            7: 'julho',
            8: 'agosto',
            9: 'setembro',
            10: 'outubro',
            11: 'novembro',
            12: 'dezembro'
        }

    def criar_carga(self):
        carga_mensal = self.documento.add_paragraph()
        carga_mensal.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        carga_mensal_run = carga_mensal.add_run(f"Carga Suplementar {self.dict_mes[self.mes_pagamento]}")
        carga_mensal_run.font.size = Pt(12)

        table_carga = self.documento.add_table(rows=carga_mensal.shape[0] + 1, cols=carga_mensal.shape[1])
        table_carga.style = 'Table Grid'  # Add borders to the table

        # Set alignment for each cell in the table
        for row in table_carga.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.documento.save(f'carga/Carga Suplementar{self.dict_mes[self.mes_pagamento]}-{self.ano_atual}')