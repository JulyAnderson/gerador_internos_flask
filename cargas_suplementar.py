from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from datetime import datetime

#Adaptação para uso com pyinstaller
import sys
import os

class CargaGenerator:
    def __init__(self, data_handler,mes_pagamento):
        self.data_handler = data_handler
        self.mes_pagamento = mes_pagamento
        self.ano_atual = datetime.now().year
        self.dict_mes = {
            1: 'janeiro',
            2: 'fevereiro',
            3: 'março',
            4: 'abril',
            5: 'maio',
            6: 'junho',
            7: 'julho',
            8: 'agosto',
            9: 'setembro',
            10: 'outubro',
            11: 'novembro',
            12: 'dezembro'
        }
    
    def criar_documento_carga(self):
        #self.documento = Document('modelo.docx')
        #cria o modelo para funcionar com o pyinstaller 
        if getattr(sys, 'frozen', False):
            # Caminho do arquivo dentro do pacote PyInstaller
            bundle_dir = sys._MEIPASS
        else:
            # Caminho do arquivo em desenvolvimento
            bundle_dir = os.path.abspath(os.path.dirname(__file__))

        modelo_path = os.path.join(bundle_dir, 'modelo.docx')
        self.documento = Document(modelo_path)

    def cabecalho(self):
        carga_mensal = self.documento.add_paragraph()
        carga_mensal.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        carga_mensal_run = carga_mensal.add_run(f"Carga Suplementar {self.dict_mes[self.mes_pagamento]}--{self.ano_atual}")
        carga_mensal_run.font.size = Pt(12)

    def tabela_carga(self, carga_mes):
        # Cria uma tabela vazia no documento
        tabela = self.documento.add_table(rows=1, cols=len(carga_mes.columns))
        tabela.style = 'Table Grid'  # Add borders to the table

        # Adiciona os cabeçalhos das colunas à primeira linha da tabela
        for i, coluna in enumerate(carga_mes.columns):
            celula = tabela.cell(0, i)
            paragrafo = celula.paragraphs[0]
            paragrafo.text = coluna
            paragrafo.runs[0].bold = True

        # Adiciona os dados do DataFrame às células da tabela
        for i, linha in carga_mes.iterrows():
            nova_linha = tabela.add_row().cells
            for j, valor in enumerate(linha):
                paragrafo = nova_linha[j].paragraphs[0]
                paragrafo.text = str(valor)

    def criar_carga(self,carga_mes):
        self.criar_documento_carga()
        self.cabecalho()
        self.tabela_carga(carga_mes)
        self.documento.save(f'carga/Carga Suplementar{self.dict_mes[self.mes_pagamento]}-{self.ano_atual}.docx')
